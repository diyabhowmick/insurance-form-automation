"""
DOCX Handler Module
Handles loading, parsing, and filling Word document templates
Supports multiple placeholder formats: [FIELD], {{FIELD}}, {FIELD}
"""

from docx import Document
from typing import List, Dict, Set
import re
import io

def load_template(docx_file) -> Document:
    """
    Load a Word document template.
    
    Args:
        docx_file: Uploaded .docx file or file path
        
    Returns:
        Document: python-docx Document object
    """
    try:
        if hasattr(docx_file, 'read'):
            # It's a file-like object (uploaded file)
            doc = Document(docx_file)
            docx_file.seek(0)  # Reset for potential reuse
        else:
            # It's a file path
            doc = Document(docx_file)
        return doc
    except Exception as e:
        raise Exception(f"Error loading template: {str(e)}")


def find_placeholders(doc: Document) -> List[str]:
    """
    Find all placeholders in the document.
    Supports formats: {{placeholder}}, [PLACEHOLDER], {placeholder}
    
    Args:
        doc: python-docx Document object
        
    Returns:
        list: List of unique placeholder names
    """
    placeholders = set()
    # Multiple placeholder patterns
    patterns = [
        r'\{\{([^}]+)\}\}',  # {{FIELD_NAME}}
        r'\[([A-Z_]+)\]',     # [FIELD_NAME]
        r'\{([A-Z_]+)\}'      # {FIELD_NAME}
    ]
    
    def extract_from_text(text):
        for pattern in patterns:
            matches = re.findall(pattern, text)
            placeholders.update(matches)
    
    # Search in paragraphs
    for paragraph in doc.paragraphs:
        extract_from_text(paragraph.text)
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    extract_from_text(paragraph.text)
    
    # Search in headers and footers
    for section in doc.sections:
        # Header
        header = section.header
        for paragraph in header.paragraphs:
            extract_from_text(paragraph.text)
        
        # Footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            extract_from_text(paragraph.text)
    
    return sorted(list(placeholders))


def replace_placeholders(doc: Document, data: Dict[str, str]) -> Document:
    """
    Replace placeholders in the document with actual values.
    Supports formats: {{placeholder}}, [PLACEHOLDER], {placeholder}
    
    Args:
        doc: python-docx Document object
        data: Dictionary mapping placeholder names to values
        
    Returns:
        Document: Modified document
    """
    
    def replace_in_paragraph(paragraph, placeholder, value):
        """Helper function to replace text in a paragraph while preserving formatting."""
        # Try different placeholder formats
        formats = [
            f"{{{{{placeholder}}}}}",  # {{FIELD_NAME}}
            f"[{placeholder}]",         # [FIELD_NAME]
            f"{{{placeholder}}}"        # {FIELD_NAME}
        ]
        
        for fmt in formats:
            if fmt in paragraph.text:
                full_text = paragraph.text
                new_text = full_text.replace(fmt, str(value))
                
                # Clear existing runs and add new text
                for run in paragraph.runs:
                    run.text = ''
                
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
                break  # Stop after first match
    
    # Replace in each placeholder
    for placeholder, value in data.items():
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, placeholder, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, placeholder, value)
        
        # Replace in headers and footers
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                replace_in_paragraph(paragraph, placeholder, value)
            
            # Footer
            for paragraph in section.footer.paragraphs:
                replace_in_paragraph(paragraph, placeholder, value)
    
    return doc


def save_document_to_bytes(doc: Document) -> bytes:
    """
    Save document to bytes for download.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        bytes: Document as bytes
    """
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def get_template_info(doc: Document) -> Dict:
    """
    Get information about the template.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        dict: Template information
    """
    return {
        'paragraph_count': len(doc.paragraphs),
        'table_count': len(doc.tables),
        'section_count': len(doc.sections),
        'placeholders': find_placeholders(doc)
    }